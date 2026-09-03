"""
Microsoft Word (.docx) Document Generator.
Places cropped and normalized cards and document sheets into ready-to-print Word files
with exact physical dimensions dynamically calculated from the active FormatProfile.
"""
from pathlib import Path
from typing import Union, List, Optional
import tempfile
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from app.core.config import PrintConfig
from app.core.profiles import FormatProfile, CARD_PROFILE, LONG_FORM_PROFILE
from app.core.models import CardPair, SheetQueue, ProcessedImage
from app.utils.logger import get_logger

logger = get_logger("docx_export")


class DocxExporter:
    """Exports processed documents and cards to Microsoft Word (.docx)."""

    @classmethod
    def export_card_pair(
        cls,
        card_pair: CardPair,
        output_path: Union[str, Path],
        layout: str = "auto",  # "auto", "side_by_side", "stacked"
        custom_width_cm: Optional[float] = None,
        custom_height_cm: Optional[float] = None,
    ) -> bool:
        """
        Exports Front and Back cards to a print-ready Microsoft Word (.docx) document.
        Always uses standard A4 portrait page layout (210 x 297 mm).
        Supports custom width/height adjustments and auto layout (Stacked for Long Form,
        Side-by-side for Standard Cards). Zero text or labels rendered on the page.
        """
        out_file = Path(output_path)
        out_file.parent.mkdir(parents=True, exist_ok=True)

        try:
            doc = Document()
            profile = card_pair.format_profile or CARD_PROFILE

            # Standard A4 Portrait Page Setup (210 mm x 297 mm)
            section = doc.sections[0]
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.2)
            section.right_margin = Cm(1.2)

            usable_w_cm = 21.0 - 2.4  # 18.6 cm

            with tempfile.TemporaryDirectory() as tmpdir:
                tmp_path = Path(tmpdir)

                front_img = card_pair.front.final_image if card_pair.front is not None else None
                back_img = card_pair.back.final_image if card_pair.back is not None else None

                front_tmp = None
                back_tmp = None

                if front_img is not None:
                    front_tmp = tmp_path / "front_card.png"
                    cv2.imwrite(str(front_tmp), front_img)

                if back_img is not None:
                    back_tmp = tmp_path / "back_card.png"
                    cv2.imwrite(str(back_tmp), back_img)

                is_long_form = (profile.id == "long_form" or profile.aspect_ratio >= 2.0)
                ref_img = front_img if front_img is not None else back_img
                is_portrait = False
                img_aspect = profile.aspect_ratio
                if ref_img is not None:
                    ih, iw = ref_img.shape[:2]
                    is_portrait = (ih > iw)
                    img_aspect = iw / max(ih, 1)

                # Determine effective layout:
                # - Portrait Long Form (8.5 x 21 cm): fits side-by-side naturally (8.5 + 8.5 = 17 cm <= 18.6 cm usable)
                # - Landscape Long Form (21 x 8.5 cm): stacked fits full page width (~18.5 cm wide each)
                # - Standard Cards: side-by-side
                if layout == "auto":
                    if is_long_form:
                        effective_layout = "side_by_side" if is_portrait else "stacked"
                    else:
                        effective_layout = "side_by_side"
                else:
                    effective_layout = layout

                # Determine dimensions strictly preserving the image's native aspect ratio
                if custom_width_cm is not None and custom_width_cm > 0:
                    card_w = Cm(custom_width_cm)
                    if custom_height_cm is not None and custom_height_cm > 0:
                        card_h = Cm(custom_height_cm)
                    else:
                        card_h = Cm(round(custom_width_cm / max(img_aspect, 1e-3), 2))
                else:
                    # Default physical dimensions
                    if is_long_form:
                        if is_portrait:
                            # Vertical Aadhaar letter cut-slip (85 x 210 mm)
                            card_w = Cm(8.5)
                            card_h = Cm(round(8.5 / max(img_aspect, 1e-3), 2))
                        else:
                            # Horizontal Aadhaar letter cut-slip (210 x 85 mm)
                            if effective_layout == "stacked":
                                card_w = Cm(18.5)
                                card_h = Cm(round(18.5 / max(img_aspect, 1e-3), 2))
                            else:
                                card_w = Cm(9.1)
                                card_h = Cm(round(9.1 / max(img_aspect, 1e-3), 2))
                    elif is_portrait:
                        card_w = Cm(min(profile.width_mm, profile.height_mm) / 10.0)
                        card_h = Cm(max(profile.width_mm, profile.height_mm) / 10.0)
                    else:
                        card_w = Cm(max(profile.width_mm, profile.height_mm) / 10.0)
                        card_h = Cm(min(profile.width_mm, profile.height_mm) / 10.0)

                if front_tmp and back_tmp:
                    if effective_layout == "side_by_side":
                        # 2-column borderless table
                        table = doc.add_table(rows=1, cols=2)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.autofit = False

                        cell_left = table.cell(0, 0)
                        cell_left.width = Cm(usable_w_cm / 2.0)
                        p_left = cell_left.paragraphs[0]
                        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_left.paragraph_format.space_before = Pt(0)
                        p_left.paragraph_format.space_after = Pt(0)
                        p_left.add_run().add_picture(str(front_tmp), width=card_w, height=card_h)

                        cell_right = table.cell(0, 1)
                        cell_right.width = Cm(usable_w_cm / 2.0)
                        p_right = cell_right.paragraphs[0]
                        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_right.paragraph_format.space_before = Pt(0)
                        p_right.paragraph_format.space_after = Pt(0)
                        p_right.add_run().add_picture(str(back_tmp), width=card_w, height=card_h)
                    else:
                        # Stacked layout (Front top, Back bottom) - Ideal for Long Form Aadhaar
                        p_top = doc.add_paragraph()
                        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_top.paragraph_format.space_before = Pt(0)
                        p_top.paragraph_format.space_after = Pt(14)
                        p_top.add_run().add_picture(str(front_tmp), width=card_w, height=card_h)

                        p_bot = doc.add_paragraph()
                        p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_bot.paragraph_format.space_before = Pt(0)
                        p_bot.paragraph_format.space_after = Pt(0)
                        p_bot.add_run().add_picture(str(back_tmp), width=card_w, height=card_h)

                elif front_tmp or back_tmp:
                    # Single side present
                    active_tmp = front_tmp or back_tmp
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.add_run().add_picture(str(active_tmp), width=card_w, height=card_h)

            doc.save(str(out_file))
            logger.info(f"Successfully exported Word document to {out_file}")
            return True

        except Exception as e:
            logger.error(f"Failed to export card pair to Word docx: {e}", exc_info=True)
            return False

    @classmethod
    def export_sheet_queue(
        cls,
        queue: SheetQueue,
        output_path: Union[str, Path]
    ) -> bool:
        """
        Exports multi-page sheet documents to Word (.docx), placing each page on a separate page.
        """
        out_file = Path(output_path)
        out_file.parent.mkdir(parents=True, exist_ok=True)

        if not queue.pages:
            logger.warning("No pages in sheet queue to export.")
            return False

        try:
            doc = Document()

            for section in doc.sections:
                section.top_margin = Inches(PrintConfig.DOCX_MARGIN_INCH)
                section.bottom_margin = Inches(PrintConfig.DOCX_MARGIN_INCH)
                section.left_margin = Inches(PrintConfig.DOCX_MARGIN_INCH)
                section.right_margin = Inches(PrintConfig.DOCX_MARGIN_INCH)

            with tempfile.TemporaryDirectory() as tmpdir:
                tmp_path = Path(tmpdir)

                for idx, page in enumerate(queue.pages):
                    img = page.final_image
                    if img is None:
                        continue

                    page_tmp = tmp_path / f"page_{idx + 1}.png"
                    cv2.imwrite(str(page_tmp), img)

                    h, w = img.shape[:2]
                    is_landscape = w > h

                    if is_landscape:
                        target_width = Inches(7.5)
                    else:
                        target_width = Inches(7.0)

                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(page_tmp), width=target_width)

                    if idx < len(queue.pages) - 1:
                        doc.add_page_break()

            doc.save(str(out_file))
            logger.info(f"Successfully exported Sheet document to {out_file}")
            return True

        except Exception as e:
            logger.error(f"Failed to export sheets to Word docx: {e}", exc_info=True)
            return False

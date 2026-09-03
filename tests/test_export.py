"""
Unit tests for Word (.docx), PDF, and Image Exporters.
"""
from pathlib import Path
import tempfile
import pytest

from docx import Document
from app.core.models import EnhancementMode
from app.processing.card_processor import CardProcessor
from app.processing.sheet_processor import SheetProcessor
from app.export.docx_export import DocxExporter
from app.export.pdf_export import PdfExporter
from app.export.image_export import ImageExporter


@pytest.fixture
def test_images_dir():
    return Path(__file__).resolve().parent.parent / "test_images"


def test_docx_export_card_pair(test_images_dir):
    cp = CardProcessor()
    f_path = test_images_dir / "cards" / "card_01_straight_front.jpg"
    b_path = test_images_dir / "cards" / "card_02_rotated_perspective_back.jpg"
    pair = cp.process_pair(f_path, b_path, enhancement_mode=EnhancementMode.DOCUMENT_CRISP)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "test_cards.docx"
        success = DocxExporter.export_card_pair(pair, docx_path)
        assert success
        assert docx_path.exists()
        assert docx_path.stat().st_size > 1000

        # Read back document with python-docx
        doc = Document(str(docx_path))
        # Should have table with 2 columns
        assert len(doc.tables) >= 1
        assert len(doc.tables[0].columns) == 2


def test_docx_export_sheet_queue(test_images_dir):
    sp = SheetProcessor()
    s1 = test_images_dir / "sheets" / "sheet_01_invoice_p1.jpg"
    s2 = test_images_dir / "sheets" / "sheet_02_invoice_p2.jpg"
    queue = sp.process_queue([s1, s2])

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "test_sheets.docx"
        success = DocxExporter.export_sheet_queue(queue, docx_path)
        assert success
        assert docx_path.exists()
        assert docx_path.stat().st_size > 1000


def test_docx_export_long_form(test_images_dir):
    cp = CardProcessor()
    f_path = test_images_dir / "real_samples" / "WhatsApp Image 2026-09-01 at 9.36.55 PM (2).jpeg"
    b_path = test_images_dir / "real_samples" / "WhatsApp Image 2026-09-01 at 9.36.35 PM.jpeg"
    if not f_path.exists():
        pytest.skip("Sample not found")

    from app.core.profiles import LONG_FORM_PROFILE
    pair = cp.process_pair(f_path, b_path, profile=LONG_FORM_PROFILE)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "test_long_form.docx"
        success = DocxExporter.export_card_pair(pair, docx_path)
        assert success
        assert docx_path.exists()
        assert docx_path.stat().st_size > 1000


def test_pdf_and_image_export(test_images_dir):
    cp = CardProcessor()
    f_path = test_images_dir / "cards" / "card_01_straight_front.jpg"
    b_path = test_images_dir / "cards" / "card_02_rotated_perspective_back.jpg"
    pair = cp.process_pair(f_path, b_path)

    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = Path(tmpdir) / "test_card.pdf"
        assert PdfExporter.export_card_pair(pair, pdf_path)
        assert pdf_path.exists()

        images = ImageExporter.export_card_pair(pair, tmpdir, format_ext=".png")
        assert len(images) == 3  # duplex side-by-side composite, front, back
        assert all(p.exists() for p in images)
